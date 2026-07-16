"""
rag_claims_kb.py
----------------
Claims Denial KB — RAG: LLM Extraction -> BM25 -> LLM Matching -> Validation.

Universal Knowledge Base Loader:
  Supports: Excel (.xlsx), Word (.docx), PDF (.pdf), HTML (.html), Markdown (.md)

  All file types can have EITHER:
    - A table with columns: Rule ID | Description | Results
    - Free-form text/paragraphs (parsed by structure)

Pipeline:
  1. LLM field extraction   -> parses NL question into structured JSON fields.
  2. BM25 retrieval         -> searches KB using description, returns top-K candidates.
  2.5 BM25 on Codes sheet   -> finds top-K relevant code categories using extracted CPTs.
  3. LLM field matching     -> strict rule-based match; reads Results field directly.
  4. LLM validation         -> verifies matched rule truly applies; uses Codes BM25 results.
"""

import json
import os
import re
import time
import requests
import numpy as np
import pandas as pd
from rank_bm25 import BM25Okapi
from pathlib import Path

# ─────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────
OLLAMA_BASE = "http://localhost:11434"
LLM_MODEL   = "qwen2.5:7b"

# ── Set your knowledge base file here ──
# Supported: .xlsx  .docx  .pdf  .html  .md
KB_FILE    = "D:\\RAG PROJECT\\claims_rules1.xlsx"

# ── Codes sheet stays as Excel ──
CODES_FILE = "D:\\RAG PROJECT\\claims_rules1.xlsx"

TOP_K = 2


# ═══════════════════════════════════════════════════════════════
# SHARED HELPER — results field parser
# ═══════════════════════════════════════════════════════════════
def _parse_results_field(results_text: str) -> tuple:
    """
    Splits Results column into (action, instruction).
    Results format: "Action: Adjust Claim. When Medicare denies..."
    """
    results_text = (results_text or "").strip()
    match = re.match(r"Action:\s*(.+?)\.\s*(.*)", results_text, re.DOTALL)
    if match:
        return match.group(1).strip(), match.group(2).strip()
    return results_text, results_text


def _build_kb_entry(rule_id: str, description: str, results: str) -> dict:
    """
    Standard KB entry builder — all loaders call this.
    Guarantees every entry has the same structure.
    """
    action, instruction = _parse_results_field(results)
    return {
        "id":          str(rule_id).strip(),
        "description": str(description).strip(),
        "results":     str(results).strip(),
        "action":      action,
        "instruction": instruction,
    }


def _rows_to_kb(rows: list, headers: list) -> list:
    """
    Converts a list of rows + headers into KB entries.
    Handles both structured (Rule ID/Description/Results columns)
    and free-form (combine all cells into one description).
    Used by multiple loaders to avoid duplicating this logic.
    """
    kb = []

    # Check if structured columns exist
    has_rule_id     = "Rule ID"     in headers
    has_description = "Description" in headers
    has_results     = "Results"     in headers
    is_structured   = has_rule_id and has_description and has_results

    for row_idx, row in enumerate(rows, 1):
        # Pad row to match headers length
        while len(row) < len(headers):
            row.append("")

        if is_structured:
            rule_id     = str(row[headers.index("Rule ID")]).strip()
            description = str(row[headers.index("Description")]).strip()
            results     = str(row[headers.index("Results")]).strip()

            if not rule_id or rule_id.lower() in ("rule id", "nan", "none"):
                continue

        else:
            # Free-form — combine all cells into one description
            parts = []
            for header, cell in zip(headers, row):
                val = str(cell).strip()
                if val and val.lower() not in ("nan", "none", ""):
                    parts.append(f"{header}: {val}" if header else val)

            combined = " | ".join(parts)
            if not combined.strip():
                continue

            rule_id     = f"rule_{row_idx:03d}"
            description = combined
            results     = combined  # LLM will parse action from this in step 3

        kb.append(_build_kb_entry(rule_id, description, results))

    return kb


# ═══════════════════════════════════════════════════════════════
# LOADER 1 — EXCEL (.xlsx / .xls)
# ═══════════════════════════════════════════════════════════════
def _load_from_excel(filepath: str) -> list:
    """
    Reads rules from Excel. Works with:
      - Structured: Rule ID | Description | Results columns
      - Free-form:  Any columns — all cells combined per row
    """
    print(f"[Loader] Reading Excel: {filepath}")
    df   = pd.read_excel(filepath, sheet_name=0)
    cols = list(df.columns)
    print(f"  Columns found: {cols}")

    rows = []
    for _, row in df.iterrows():
        rows.append([str(row[c]) for c in cols])

    kb = _rows_to_kb(rows, cols)
    print(f"  Extracted {len(kb)} rules from Excel")
    return kb


# ═══════════════════════════════════════════════════════════════
# LOADER 2 — WORD DOCUMENT (.docx)
# ═══════════════════════════════════════════════════════════════
def _load_from_docx(filepath: str) -> list:
    """
    Reads rules from Word. Works with:
      - Tables: header row + data rows
      - Paragraphs: groups of text separated by blank lines
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx not installed.\n"
            "Run: pip install python-docx --break-system-packages"
        )

    print(f"[Loader] Reading Word document: {filepath}")
    doc = Document(filepath)
    kb  = []

    # ── Try tables first ──
    if doc.tables:
        print(f"  Found {len(doc.tables)} table(s) — reading table rows")
        for table in doc.tables:
            if not table.rows:
                continue
            headers  = [c.text.strip() for c in table.rows[0].cells]
            data_rows = [[c.text.strip() for c in row.cells] for row in table.rows[1:]]
            kb.extend(_rows_to_kb(data_rows, headers))

    # ── Fall back to paragraphs ──
    if not kb:
        print(f"  No tables — reading paragraphs")
        current, rule_num = [], 1
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                current.append(text)
            else:
                if current:
                    combined = " ".join(current)
                    kb.append(_build_kb_entry(
                        f"rule_{rule_num:03d}", combined, combined
                    ))
                    current = []
                    rule_num += 1
        if current:
            combined = " ".join(current)
            kb.append(_build_kb_entry(f"rule_{rule_num:03d}", combined, combined))

    print(f"  Extracted {len(kb)} rules from Word")
    return kb


# ═══════════════════════════════════════════════════════════════
# LOADER 3 — PDF (.pdf)
# ═══════════════════════════════════════════════════════════════
def _load_from_pdf(filepath: str) -> list:
    """
    Reads rules from PDF. Works with:
      - Tables: extracted by pdfplumber
      - Text: paragraphs separated by blank lines
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError(
            "pdfplumber not installed.\n"
            "Run: pip install pdfplumber --break-system-packages"
        )

    print(f"[Loader] Reading PDF: {filepath}")
    kb       = []
    headers  = None
    all_rows = []
    rule_num = 1

    with pdfplumber.open(filepath) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            tables = page.extract_tables()

            if tables:
                for table in tables:
                    if not table:
                        continue
                    first_row = [str(c).strip() if c else "" for c in table[0]]

                    # Detect header row
                    is_header = any(
                        kw in " ".join(first_row).lower()
                        for kw in ["rule", "description", "results", "action", "denial"]
                    )

                    if headers is None or is_header:
                        headers   = first_row
                        data_rows = table[1:]
                    else:
                        data_rows = table

                    all_rows.extend(
                        [[str(c).strip() if c else "" for c in row]
                         for row in data_rows]
                    )

            else:
                # Text fallback
                text = page.extract_text() or ""
                current = []
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        current.append(line)
                    else:
                        if current:
                            combined = " ".join(current)
                            kb.append(_build_kb_entry(
                                f"rule_{rule_num:03d}", combined, combined
                            ))
                            current  = []
                            rule_num += 1
                if current:
                    combined = " ".join(current)
                    kb.append(_build_kb_entry(
                        f"rule_{rule_num:03d}", combined, combined
                    ))

    if all_rows:
        kb.extend(_rows_to_kb(all_rows, headers or []))

    print(f"  Extracted {len(kb)} rules from PDF")
    return kb


# ═══════════════════════════════════════════════════════════════
# LOADER 4 — HTML (.html / .htm)  ← NEW
# ═══════════════════════════════════════════════════════════════
def _load_from_html(filepath: str) -> list:
    """
    Reads rules from an HTML file. Works with:
      - <table> elements: header row + data rows
      - <p> or <li> paragraphs: each becomes one rule

    Install: pip install beautifulsoup4 --break-system-packages
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError(
            "beautifulsoup4 not installed.\n"
            "Run: pip install beautifulsoup4 --break-system-packages"
        )

    print(f"[Loader] Reading HTML: {filepath}")

    with open(filepath, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    kb = []

    # ── Try <table> elements first ──
    tables = soup.find_all("table")
    if tables:
        print(f"  Found {len(tables)} table(s) in HTML")
        for table in tables:
            rows     = table.find_all("tr")
            if not rows:
                continue

            # First row = headers
            header_cells = rows[0].find_all(["th", "td"])
            headers      = [c.get_text(strip=True) for c in header_cells]
            print(f"  HTML table headers: {headers}")

            data_rows = []
            for row in rows[1:]:
                cells = row.find_all(["th", "td"])
                data_rows.append([c.get_text(strip=True) for c in cells])

            kb.extend(_rows_to_kb(data_rows, headers))

    # ── Fall back to paragraphs and list items ──
    if not kb:
        print(f"  No tables — reading <p> and <li> elements")
        rule_num = 1
        elements = soup.find_all(["p", "li", "div"])
        for elem in elements:
            text = elem.get_text(strip=True)
            if text and len(text) > 20:  # skip tiny/empty elements
                kb.append(_build_kb_entry(
                    f"rule_{rule_num:03d}", text, text
                ))
                rule_num += 1

    print(f"  Extracted {len(kb)} rules from HTML")
    return kb


# ═══════════════════════════════════════════════════════════════
# LOADER 5 — MARKDOWN (.md)  ← NEW
# ═══════════════════════════════════════════════════════════════
def _load_from_markdown(filepath: str) -> list:
    """
    Reads rules from a Markdown file. Works with:
      - Pipe tables: | Rule ID | Description | Results |
      - Paragraphs:  text blocks separated by blank lines
      - Lists:       - rule text or * rule text

    No extra package needed — uses built-in re module.
    """
    print(f"[Loader] Reading Markdown: {filepath}")

    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    kb       = []
    rule_num = 1

    # ── Try pipe tables first ──
    # Markdown table format:
    # | Header1 | Header2 | Header3 |
    # |---------|---------|---------|
    # | val1    | val2    | val3    |

    table_pattern = re.compile(
        r"(\|.+\|)\n\|[-| :]+\|\n((?:\|.+\|\n?)+)",
        re.MULTILINE
    )

    for table_match in table_pattern.finditer(content):
        header_line = table_match.group(1)
        body_lines  = table_match.group(2).strip().split("\n")

        # Parse headers
        headers = [h.strip() for h in header_line.split("|") if h.strip()]
        print(f"  Markdown table headers: {headers}")

        # Parse data rows
        data_rows = []
        for line in body_lines:
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if cells:
                data_rows.append(cells)

        kb.extend(_rows_to_kb(data_rows, headers))

    # ── Fall back to paragraphs if no tables found ──
    if not kb:
        print(f"  No tables — reading paragraphs and list items")

        # Remove markdown formatting characters for clean text
        clean = re.sub(r"#{1,6}\s+", "", content)         # remove headings ##
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)    # remove bold **text**
        clean = re.sub(r"\*(.+?)\*",     r"\1", clean)    # remove italic *text*
        clean = re.sub(r"`(.+?)`",       r"\1", clean)    # remove code `text`

        # Split into blocks by blank lines
        blocks = re.split(r"\n\s*\n", clean)

        for block in blocks:
            block = block.strip()

            # Skip empty blocks, pure separators, heading-only blocks
            if not block or block.startswith("---") or len(block) < 15:
                continue

            # Clean list markers (- item or * item or 1. item)
            block = re.sub(r"^[\-\*\d+\.]\s+", "", block, flags=re.MULTILINE)
            block = " ".join(block.split())  # collapse whitespace

            if block:
                kb.append(_build_kb_entry(
                    f"rule_{rule_num:03d}", block, block
                ))
                rule_num += 1

    print(f"  Extracted {len(kb)} rules from Markdown")
    return kb


# ═══════════════════════════════════════════════════════════════
# UNIVERSAL LOADER — entry point
# ═══════════════════════════════════════════════════════════════
def load_rules_from_file(filepath: str) -> list:
    """
    Universal knowledge base loader.
    Detects file extension and calls the correct loader.

    Supported formats:
      .xlsx / .xls  → Excel
      .docx         → Word
      .pdf          → PDF
      .html / .htm  → HTML  (NEW)
      .md           → Markdown  (NEW)

    All loaders return the same KB entry format:
    {
      "id":          str,
      "description": str,
      "results":     str,
      "action":      str,
      "instruction": str
    }
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"Knowledge base file not found: {filepath}")

    ext = Path(filepath).suffix.lower()

    loader_map = {
        ".xlsx": _load_from_excel,
        ".xls":  _load_from_excel,
        ".docx": _load_from_docx,
        ".pdf":  _load_from_pdf,
        ".html": _load_from_html,
        ".htm":  _load_from_html,
        ".md":   _load_from_markdown,
    }

    if ext not in loader_map:
        raise ValueError(
            f"Unsupported file type: '{ext}'\n"
            f"Supported: {', '.join(sorted(loader_map.keys()))}"
        )

    print(f"\n{'='*55}")
    print(f"Loading KB: {Path(filepath).name}  [{ext}]")
    print(f"{'='*55}")

    kb = loader_map[ext](filepath)

    if not kb:
        raise ValueError(f"No rules could be extracted from: {filepath}")

    print(f"\nTotal rules loaded: {len(kb)}")
    return kb


# ═══════════════════════════════════════════════════════════════
# CODES SHEET LOADER — always Excel
# ═══════════════════════════════════════════════════════════════
def load_codes_from_excel(filepath: str) -> list:
    try:
        df = pd.read_excel(filepath, sheet_name="Codes")
    except Exception:
        print("  WARNING: No Codes sheet found — CPT category check disabled")
        return []

    codes = []
    for _, row in df.iterrows():
        code = str(row.get("Code", "")).strip()
        cpt  = str(row.get("CPT",  "")).strip()
        if code and cpt and code.lower() != "nan":
            codes.append({"code": code, "cpt": cpt})
    return codes


# ─────────────────────────────────────────────────────────────
# STARTUP — load everything
# ─────────────────────────────────────────────────────────────
KB       = load_rules_from_file(KB_FILE)
CODES_KB = load_codes_from_excel(CODES_FILE)

print(f"Loaded {len(CODES_KB)} code categories from Codes sheet")

if KB:
    print(f"\nSample rule:")
    print(f"  id          : {KB[0]['id']}")
    print(f"  action      : {KB[0]['action']}")
    print(f"  instruction : {KB[0]['instruction'][:80]}...")


# ─────────────────────────────────────────────────────────────
# BM25 INDEX — Rules
# ─────────────────────────────────────────────────────────────
BM25_DOCS = []
for entry in KB:
    text = f"{entry['description']} {entry['results']}"
    BM25_DOCS.append(text.lower().split())

bm25 = BM25Okapi(BM25_DOCS)
print(f"\nBM25 index built for {len(KB)} rules")

# ─────────────────────────────────────────────────────────────
# BM25 INDEX — Codes sheet
# ─────────────────────────────────────────────────────────────
CODES_BM25_DOCS = []
for entry in CODES_KB:
    CODES_BM25_DOCS.append(f"{entry['code']} {entry['cpt']}".lower().split())

bm25_codes = BM25Okapi(CODES_BM25_DOCS) if CODES_BM25_DOCS else None
print(f"BM25 index built for {len(CODES_KB)} code categories")


# ─────────────────────────────────────────────────────────────
# STEP 1 — LLM FIELD EXTRACTION
# ─────────────────────────────────────────────────────────────
_EXTRACT_SYSTEM = """You are a structured data extractor for a medical billing system.
Extract the following fields from the user's question and return ONLY a valid JSON object -- no markdown, no explanation.

Fields:
  - group             : e.g. "Group 1", "Group 2"          (null if not mentioned)
  - practice          : practice code e.g. "MICLAI"         (null if not mentioned)
  - insurance_company : e.g. "Medicare", "Aetna"            (null if not mentioned)
  - plan_name         : e.g. "MED", "HMO"                   (null if not mentioned)
  - cpt_codes         : list of CPT/J-codes e.g. ["99454"]  ([] if not mentioned)
  - denial_code       : numeric string e.g. "151"           (null if not mentioned)
  - remark_code       : e.g. "M25"                          (null if not mentioned)

Return ONLY the JSON object."""


def extract_fields(question: str) -> dict:
    t_start = time.time()
    resp = requests.post(
        f"{OLLAMA_BASE}/api/generate",
        json={
            "model":  LLM_MODEL,
            "system": _EXTRACT_SYSTEM,
            "prompt": f"Extract fields from:\n{question}",
            "stream": False,
            "options": {"temperature": 0.0},
        },
        timeout=600,
    )
    resp.raise_for_status()
    raw        = resp.json()["response"].strip()
    elapsed_ms = int((time.time() - t_start) * 1000)

    json_match = re.search(r"\{.*\}", raw, re.DOTALL)
    if not json_match:
        raise ValueError(f"LLM did not return valid JSON.\nRaw: {raw}")

    extracted = json.loads(json_match.group())
    print(f"Extracted fields ({elapsed_ms} ms): {json.dumps(extracted)}")
    return extracted


# ─────────────────────────────────────────────────────────────
# STEP 2 — BM25 RETRIEVAL (Rules)
# ─────────────────────────────────────────────────────────────
def retrieve_candidates(extracted: dict, k: int = TOP_K) -> list:
    query = (
        f"{extracted.get('insurance_company', '')} "
        f"{' '.join(extracted.get('cpt_codes', []))} "
        f"{extracted.get('denial_code', '')} "
        f"{extracted.get('remark_code', '')}"
    )

    tokens = query.lower().split()
    scores = bm25.get_scores(tokens)
    ranked = np.argsort(scores)[::-1][:k]

    results = []
    for idx in ranked:
        results.append((KB[idx], float(scores[idx])))

    print(f"\nBM25 top-{k} candidates:")
    for i, (entry, score) in enumerate(results, 1):
        print(f"  [{i}] score={score:.4f} id={entry['id']}  action={entry['action']}")

    return results


# ─────────────────────────────────────────────────────────────
# STEP 2.5 — BM25 RETRIEVAL (Codes sheet)
# ─────────────────────────────────────────────────────────────
def retrieve_code_candidates(extracted: dict, k: int = 2) -> list:
    if not bm25_codes or not CODES_KB:
        return []

    cpt_codes = extracted.get("cpt_codes", [])
    if not cpt_codes:
        print("\nCodes BM25: No CPT codes extracted — skipping")
        return []

    tokens = " ".join(cpt_codes).lower().split()
    scores = bm25_codes.get_scores(tokens)
    ranked = np.argsort(scores)[::-1][:k]

    results = []
    for idx in ranked:
        if scores[idx] > 0:
            results.append((CODES_KB[idx], float(scores[idx])))

    print(f"\nCodes BM25 top-{k} candidates:")
    if results:
        for i, (entry, score) in enumerate(results, 1):
            print(f"  [{i}] score={score:.4f} code={entry['code']} cpt={entry['cpt']}")
    else:
        print("  No relevant code categories found")

    return results


# ─────────────────────────────────────────────────────────────
# STEP 3 — LLM FIELD MATCHING
# ─────────────────────────────────────────────────────────────
_MATCH_SYSTEM = """You are a medical billing rules matcher. You will be given:
  1. EXTRACTED FIELDS -- a JSON object parsed from the user question.
  2. KB CANDIDATES    -- a numbered list of Knowledge Base entries.
                        Each candidate has:
                          - 'description': rule context in natural language ending with Keywords.
                          - 'results': "Action: <action>. <instruction text>"

FIELD ACCESS RULE:
  - If a field key is missing from EXTRACTED FIELDS, treat it as null.

HOW TO READ THE DESCRIPTION:
  - Description pattern:
      "This rule applies to <group>, practice <practice>, insurance: <insurance>,
       plan: <plan>. It covers CPT code(s) <cpts> with denial code <denial>
       and remark code <remark>. Category: <category>. Keywords: <keywords>"
  - Parse group, practice, insurance, plan, CPT codes, denial code, remark code
    from the description before matching.

HOW TO READ RESULTS:
  - Results pattern: "Action: <action>. <instruction>"
  - ACTION is the text after "Action:" up to the first period.
  - INSTRUCTION is everything after that first period.
  - Return both verbatim.

MATCHING RULES (apply strictly, field by field):
  - group         : "All" in description -> always matches.
                    Specific AND extracted null -> NO MATCH.
                    Specific AND extracted not null -> must match (case-insensitive).
  - practice      : Same rule as group.
  - insurance     : Same rule as group.
  - plan          : Same rule as group.
  - denial_code   : "All" in description -> always matches.
                    Specific AND extracted null -> NO MATCH.
                    Specific AND extracted not null -> must match exactly.
  - remark_code   : Same rule as denial_code.
  - cpt_codes     : "All" in description -> always matches.
                    Extracted [] or missing -> always matches.
                    Otherwise -> at least one extracted CPT must appear in description.

DECISION:
  - Check every candidate in order.
  - FIRST candidate where ALL fields match is the result.

OUTPUT FORMAT — Return ONLY valid JSON, no markdown:

If a match is found:
{
  "matched": true,
  "rule_id": "<matched rule id>",
  "action": "<action text from Results — verbatim>",
  "confidence": <0.0 to 1.0>,
  "confidence_label": "<High — safe to act | Medium — please verify | Low — manual review recommended>",
  "what_matched": ["<field: extracted value matched KB value>"],
  "reason": "<one sentence why this rule applies>",
  "instruction": "<instruction text from Results — verbatim>"
}

If NO match:
{
  "matched": false,
  "rule_id": null,
  "action": "Manual Review Required",
  "confidence": 0.0,
  "confidence_label": "Low — manual review recommended",
  "what_matched": [],
  "reason": "No rule in the top candidates matched all required fields.",
  "instruction": "No matching rule found in the KB."
}

STRICT RULES:
  - Output ONLY the JSON. Nothing before or after.
  - Copy action and instruction verbatim from Results.
  - Do not hallucinate values."""


def _format_candidates(candidates: list) -> str:
    lines = []
    for i, (entry, sim) in enumerate(candidates, 1):
        lines.append(
            f"[{i}] id={entry['id']}  score={sim:.4f}\n"
            f"    description: {entry['description']!r}\n"
            f"    results: {entry['results']!r}"
        )
    return "\n\n".join(lines)


def llm_match(extracted: dict, candidates: list) -> tuple:
    prompt = (
        f"EXTRACTED FIELDS:\n{json.dumps(extracted, indent=2)}\n\n"
        f"KB CANDIDATES:\n{_format_candidates(candidates)}"
    )

    t_start = time.time()
    resp = requests.post(
        f"{OLLAMA_BASE}/api/generate",
        json={
            "model":  LLM_MODEL,
            "system": _MATCH_SYSTEM,
            "prompt": prompt,
            "stream": False,
            "options": {"temperature": 0.0, "top_p": 0.9},
        },
        timeout=600,
    )
    resp.raise_for_status()

    elapsed_ms = int((time.time() - t_start) * 1000)
    answer     = resp.json().get("response", "")
    answer     = re.sub(r"<think>.*?</think>", "", answer, flags=re.DOTALL).strip()

    json_match = re.search(r"\{.*\}", answer, re.DOTALL)
    if not json_match:
        return ({
            "matched": False, "action": "Manual Review Required",
            "confidence": 0.0, "confidence_label": "Low — manual review recommended",
            "what_matched": [], "reason": "LLM did not return valid JSON.",
            "instruction": "No matching rule found in the KB.", "rule_id": None,
        }, elapsed_ms)

    try:
        return json.loads(json_match.group()), elapsed_ms
    except json.JSONDecodeError:
        return ({
            "matched": False, "action": "Manual Review Required",
            "confidence": 0.0, "confidence_label": "Low — manual review recommended",
            "what_matched": [], "reason": "LLM returned malformed JSON.",
            "instruction": "No matching rule found in the KB.", "rule_id": None,
        }, elapsed_ms)


# ─────────────────────────────────────────────────────────────
# STEP 4 — LLM VALIDATION
# ─────────────────────────────────────────────────────────────
_VALIDATE_SYSTEM = """You are a medical billing rule validator.

You will be given:
1. MATCHED RULE INSTRUCTION — instruction text from Results column
2. ACTION — action decided by the matcher
3. ORIGINAL INPUT — the user's original question
4. CPT CODES REFERENCE — top CPT categories from Codes sheet via BM25.
   Use to identify if a CPT belongs to a category like E&M.
   If not found in sheet, use your own medical billing knowledge.
   Always state source — "from sheet" or "from LLM knowledge".

Task:
- Understand the matched rule instruction and action
- Compare against the ORIGINAL INPUT
- Decide whether the rule truly applies

Rules:
- Use only information explicitly present in the input
- Do not assume missing values
- Mark unsupported conditions as failed
- Use INSUFFICIENT_DATA if critical info is missing
- Be strict — do not force a match
- If rule does not match, applicable_action must be null

Return ONLY valid JSON:
{
  "match_status": "MATCH | NO_MATCH | INSUFFICIENT_DATA",
  "is_match": true,
  "reason": "short clear summary",
  "rule_understanding": ["condition 1", "condition 2"],
  "conditions_met": ["..."],
  "conditions_failed": ["..."],
  "conditions_unverifiable": ["..."],
  "applicable_action": "text or null",
  "cpt_category_check": {
    "was_needed": true,
    "results": [
      {"cpt": "99214", "category": "E&M", "source": "from sheet"}
    ]
  }
}"""


def validate_instruction_against_input(
    instruction_text: str,
    action_text: str,
    question: str,
    code_candidates: list,
) -> tuple:
    clean_instruction = (instruction_text or "").strip()
    clean_action      = (action_text or "").strip()

    if not clean_instruction or clean_instruction == "No matching rule found in the KB.":
        return ({
            "match_status": "NO_MATCH", "is_match": False,
            "reason": "No matching KB instruction available.",
            "rule_understanding": [], "conditions_met": [],
            "conditions_failed": ["No matched KB instruction available"],
            "conditions_unverifiable": [], "applicable_action": None,
            "cpt_category_check": {"was_needed": False, "results": []},
            "validation_performed": True,
        }, 0)

    codes_text = (
        json.dumps(
            [{"code": e["code"], "cpt": e["cpt"]} for e, _ in code_candidates],
            indent=2
        )
        if code_candidates
        else "No relevant code categories found."
    )

    prompt = (
        f"MATCHED RULE INSTRUCTION:\n{clean_instruction}\n\n"
        f"ACTION:\n{clean_action}\n\n"
        f"ORIGINAL INPUT:\n{question}\n\n"
        f"CPT CODES REFERENCE (top matches from Codes sheet):\n{codes_text}\n"
        f"State whether each CPT category was 'from sheet' or 'from LLM knowledge'."
    )

    t_start = time.time()
    resp = requests.post(
        f"{OLLAMA_BASE}/api/generate",
        json={
            "model":  LLM_MODEL,
            "system": _VALIDATE_SYSTEM,
            "prompt": prompt,
            "stream": False,
            "options": {"temperature": 0.0, "top_p": 0.9},
        },
        timeout=600,
    )
    resp.raise_for_status()

    raw        = resp.json()["response"].strip()
    elapsed_ms = int((time.time() - t_start) * 1000)
    raw        = re.sub(r"<think>.*?</think>", "", raw, flags=re.DOTALL).strip()

    json_match = re.search(r"\{.*\}", raw, re.DOTALL)
    if not json_match:
        return ({
            "match_status": "INSUFFICIENT_DATA", "is_match": False,
            "reason": "Validator returned invalid JSON.",
            "rule_understanding": [], "conditions_met": [], "conditions_failed": [],
            "conditions_unverifiable": ["Validator output not valid JSON"],
            "applicable_action": None,
            "cpt_category_check": {"was_needed": False, "results": []},
            "validation_performed": True,
        }, elapsed_ms)

    try:
        validation = json.loads(json_match.group())
    except Exception as exc:
        validation = {
            "match_status": "INSUFFICIENT_DATA", "is_match": False,
            "reason": f"JSON parse error: {exc}",
            "rule_understanding": [], "conditions_met": [], "conditions_failed": [],
            "conditions_unverifiable": ["JSON parsing failed"],
            "applicable_action": None,
            "cpt_category_check": {"was_needed": False, "results": []},
        }

    validation["validation_performed"] = True
    return validation, elapsed_ms


# ─────────────────────────────────────────────────────────────
# MAIN ASK FUNCTION
# ─────────────────────────────────────────────────────────────
def ask(question: str) -> dict:

    try:
        extracted = extract_fields(question)
    except Exception as exc:
        return {
            "matched": False, "rule_id": None,
            "action": "Field extraction failed", "confidence": 0.0,
            "confidence_label": "Low — manual review recommended",
            "what_matched": [], "reason": str(exc),
            "instruction": "No matching rule found in the KB.",
            "extracted": None, "candidates": [], "code_candidates": [],
            "response_time_ms": 0, "validation": {},
        }

    candidates      = retrieve_candidates(extracted)
    code_candidates = retrieve_code_candidates(extracted)
    match_result, elapsed_ms = llm_match(extracted, candidates)

    validation, _ = validate_instruction_against_input(
        instruction_text=match_result.get("instruction", ""),
        action_text=match_result.get("action", ""),
        question=question,
        code_candidates=code_candidates,
    )

    return {
        "matched":          match_result.get("matched", False),
        "rule_id":          match_result.get("rule_id"),
        "action":           match_result.get("action", "Manual Review Required"),
        "confidence":       match_result.get("confidence", 0.0),
        "confidence_label": match_result.get("confidence_label", "Low — manual review recommended"),
        "what_matched":     match_result.get("what_matched", []),
        "reason":           match_result.get("reason", ""),
        "instruction":      match_result.get("instruction", "No matching rule found in the KB."),
        "extracted":        extracted,
        "candidates":       [e["id"] for e, _ in candidates],
        "code_candidates":  [e["code"] for e, _ in code_candidates],
        "response_time_ms": elapsed_ms,
        "validation":       validation,
    }


# ─────────────────────────────────────────────────────────────
# INTERACTIVE CLI
# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    sample_questions = [
        "Get the comments for group 1 MICLAI practice with Medicare insurance and MED plan with CPT 99454 and denial code 151",
        "What should I do for a claim denied with code 29 and no proof of timely filing?",
        "J2356 denied with code 96 for AAASC practice Group 1 all insurance all plan",
        "CPT 95251 denied as inclusive by Medicare for MICLAI Group 1 MED plan with denial 97",
        "Lab code 83036 denied with denial 50 remark M25 for AZEN Group 2 all insurance",
        "get the comments for the denial code 29",
    ]

    print("\n" + "=" * 60)
    print("  Claims Denial KB — RAG  (Universal KB Loader)")
    print(f"  KB File  : {Path(KB_FILE).name}  [{Path(KB_FILE).suffix}]")
    print("  Supported: .xlsx  .docx  .pdf  .html  .md")
    print("  Step 1   : LLM field extraction")
    print("  Step 2   : BM25 retrieval — Rules KB")
    print("  Step 2.5 : BM25 retrieval — Codes sheet")
    print("  Step 3   : LLM field matching")
    print("  Step 4   : LLM validation")
    print("  Type 'quit' to exit")
    print("=" * 60)
    print("\nSample questions (type 1–6):")
    for i, q in enumerate(sample_questions, 1):
        print(f"  {i}. {q}")

    while True:
        question = input("\nYour question: ").strip()
        if not question:
            continue
        if question.lower() in ("quit", "exit"):
            print("Goodbye!")
            break
        if question.isdigit() and 1 <= int(question) <= len(sample_questions):
            question = sample_questions[int(question) - 1]
            print(f"  -> {question}")

        result = ask(question)

        print(f"\n{'─' * 60}")
        print(f"BM25 candidates  : {result['candidates']}")
        print(f"Code candidates  : {result.get('code_candidates', [])}")
        print(f"Extracted fields : {json.dumps(result.get('extracted'), indent=2)}")
        print(f"Time             : {result.get('response_time_ms', 'N/A')} ms")
        print(f"\nMatched          : {result.get('matched')}")
        print(f"Rule ID          : {result.get('rule_id')}")
        print(f"Action           : {result.get('action')}")
        print(f"Confidence       : {result.get('confidence')} — {result.get('confidence_label')}")
        print(f"\nWhat matched:")
        for line in result.get("what_matched", []):
            print(f"  - {line}")
        print(f"\nReason           : {result.get('reason')}")
        print(f"\nInstruction      : {result.get('instruction')}")

        v = result.get("validation", {})
        print(f"\n{'─' * 60}")
        print(f"VALIDATION RESULT")
        print(f"{'─' * 60}")
        print(f"Status           : {v.get('match_status')}")
        print(f"Is Match         : {v.get('is_match')}")
        print(f"Reason           : {v.get('reason')}")
        print(f"\nRule understood as:")
        for line in v.get("rule_understanding", []):
            print(f"  - {line}")
        print(f"\nConditions met:")
        for line in v.get("conditions_met", []):
            print(f"  ✓ {line}")
        print(f"\nConditions failed:")
        for line in v.get("conditions_failed", []):
            print(f"  ✗ {line}")
        print(f"\nApplicable Action: {v.get('applicable_action')}")

        cpt_check = v.get("cpt_category_check", {})
        if cpt_check.get("was_needed"):
            print(f"\nCPT Category Check : Required")
            for r in cpt_check.get("results", []):
                print(f"  CPT {r.get('cpt')} → {r.get('category')} [{r.get('source')}]")
        else:
            print(f"\nCPT Category Check : Not required for this rule")
        print("─" * 60)
